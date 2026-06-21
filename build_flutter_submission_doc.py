from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("mobile_app_lab_study_buddy_submission.docx")
SCREENSHOT_DIR = Path("screenshots")
GITHUB_URL = "https://github.com/gandhiparas90/m-app"


def set_run_font(run, size=11, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading(doc, text, size=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_link_line(doc, label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True)
    url_run = p.add_run(url)
    set_run_font(url_run)
    url_run.font.underline = True
    return p


def add_screenshot(doc, label, image_name):
    image_path = SCREENSHOT_DIR / image_name
    if not image_path.exists():
        raise FileNotFoundError(f"Missing screenshot: {image_path}")

    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(10)
    caption.paragraph_format.space_after = Pt(4)
    caption_run = caption.add_run(label)
    set_run_font(caption_run, bold=True)

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = image_paragraph.add_run().add_picture(str(image_path), width=Inches(2.55))
    picture._inline.docPr.set("descr", label)
    picture._inline.docPr.set("title", label)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Mobile App Lab: UI Design in Flutter")
    set_run_font(run, size=26)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Study Buddy Flutter App")
    set_run_font(subtitle_run, size=12, bold=True)

    add_heading(doc, "GitHub Repository")
    add_link_line(doc, "Project link", GITHUB_URL)

    add_heading(doc, "Brief Explanation")
    add_paragraph(
        doc,
        "My Flutter app is called Study Buddy. It uses both StatelessWidget and "
        "StatefulWidget to meet the lab requirements. StudyBuddyApp and HomeScreen "
        "are stateless widgets because they display fixed interface elements such "
        "as the title, icon, short description, and navigation button. These widgets "
        "do not manage changing data internally. The home screen uses Navigator.push "
        "to move to the second screen.",
    )
    add_paragraph(
        doc,
        "The NotesScreen is a StatefulWidget because it manages dynamic content. "
        "It includes a TextField controlled by a TextEditingController, a Save Note "
        "button, and a Clear Note button. When the user saves or clears a note, "
        "setState() updates the _noteText value and rebuilds the screen so the "
        "displayed note changes immediately. This demonstrates how stateful widgets "
        "handle user interaction and changing UI data.",
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Screenshots")
    add_paragraph(
        doc,
        "The screenshots below show the home screen, the empty notes screen, a saved "
        "study note, and the cleared note state.",
    )

    screenshots = [
        ("Home screen with the Study Buddy title and navigation button", "study_buddy_home.png"),
        ("Notes screen before entering text", "study_buddy_notes_empty.png"),
        ("Notes screen after saving a study note", "study_buddy_notes_saved.png"),
        ("Notes screen after clearing the note", "study_buddy_notes_cleared.png"),
    ]
    for index, (label, image_name) in enumerate(screenshots):
        if index > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_screenshot(doc, label, image_name)

    add_heading(doc, "Screenshot Files")
    add_paragraph(
        doc,
        "These image files are also available in the repository under the screenshots "
        "folder.",
    )

    doc.save(OUT)
    print(f"wrote {OUT.resolve()}")


if __name__ == "__main__":
    build_doc()
