from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("gesture_controlled_interface_submission.docx")
SCREENSHOT_DIR = Path("screenshots")
GITHUB_URL = "https://github.com/gandhiparas90/m-app"


REFLECTION = (
    "For my gesture-controlled interface lab, I updated my Flutter app into a "
    "Gesture Studio experience. I implemented three gestures: tap, long press, "
    "and horizontal swipe. The tap gesture changes the main card color and "
    "updates the status message. This improves user experience because tapping "
    "is one of the most familiar mobile gestures and gives users a direct way "
    "to trigger immediate feedback. The long press gesture toggles a focus mode "
    "and displays a short SnackBar message. This gives the app a secondary "
    "interaction that does not require another permanent button on the screen. "
    "The horizontal swipe gesture moves between gesture cards, allowing the user "
    "to browse interaction examples in a way that feels natural on a touchscreen.\n\n"
    "These gestures support the Chapter 5 idea that mobile interfaces should "
    "connect touch input to clear, observable responses. Tap is highly "
    "discoverable, while long press and swipe can improve efficiency when they "
    "are supported with visible feedback. To make the gestures easier to use, I "
    "included a status panel that explains the result of each action. I also "
    "kept a reset button visible so users can return to the starting state if "
    "they want to try the gestures again.\n\n"
    "One challenge I encountered was that the swipe gesture could be triggered "
    "too easily during testing. That created a usability concern because users "
    "might accidentally move to another card when they only meant to touch the "
    "screen. I resolved this by adding a small velocity threshold. If the swipe "
    "is too light, the app does not change cards and instead displays a message "
    "asking the user to try a stronger swipe. I also refined the layout after "
    "testing because the gesture card could overflow on smaller screens. I "
    "fixed that by reducing the card height and making the page scrollable."
)


REFERENCE = (
    "Payne, R. (2019). Chapter 5: Mobile user experience and gestures. "
    "Course materials."
)


def set_run_font(run, size=11, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text, *, bold=False):
    for index, part in enumerate(text.split("\n\n")):
        if index:
            doc.add_paragraph()
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(part)
        set_run_font(run, bold=bold)


def add_heading(doc, text, size=16):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=size)


def add_link_line(doc, label, url):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True)
    url_run = paragraph.add_run(url)
    set_run_font(url_run)
    url_run.font.underline = True


def add_screenshot(doc, label, image_name):
    image_path = SCREENSHOT_DIR / image_name
    if not image_path.exists():
        raise FileNotFoundError(f"Missing screenshot: {image_path}")

    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(10)
    caption.paragraph_format.space_after = Pt(4)
    caption_run = caption.add_run(label)
    set_run_font(caption_run, bold=True)

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = image_paragraph.add_run().add_picture(str(image_path), width=Inches(2.55))
    picture._inline.docPr.set("descr", label)
    picture._inline.docPr.set("title", label)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Gesture-Controlled Interface")
    set_run_font(title_run, size=26)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Gesture Studio Flutter App")
    set_run_font(subtitle_run, size=12, bold=True)

    add_heading(doc, "GitHub Repository")
    add_link_line(doc, "Project link", GITHUB_URL)

    add_heading(doc, "Reflection")
    add_paragraph(doc, REFLECTION)

    add_heading(doc, "Reference")
    add_paragraph(doc, REFERENCE)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Screenshots")
    add_paragraph(
        doc,
        "The screenshots below show the home screen, the initial gesture screen, "
        "and the observable results of tap, long press, and swipe interactions.",
    )

    screenshots = [
        ("Home screen", "gesture_lab_home.png"),
        ("Initial gesture screen", "gesture_lab_initial.png"),
        ("Tap gesture changes the card color and status message", "gesture_lab_tap.png"),
        ("Long press toggles focus mode and displays feedback", "gesture_lab_long_press.png"),
        ("Swipe moves to the next gesture card", "gesture_lab_swipe.png"),
    ]

    for index, (label, image_name) in enumerate(screenshots):
        if index > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_screenshot(doc, label, image_name)

    doc.save(OUT)
    print(f"wrote {OUT.resolve()}")


if __name__ == "__main__":
    build_doc()
